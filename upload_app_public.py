import streamlit as st
from docx import Document
import pandas as pd
import io
from openai import OpenAI
import os
from prompts import PROMPTS

# Initialize OpenAI client
#API_KEY = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else os.getenv("OPENAI_API_KEY", "")
API_KEY = ''
client = OpenAI(api_key=API_KEY)

def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())
    return '\n'.join(text)

def process_file_content(uploaded_file):
    try:
        file_type = uploaded_file.type
        if file_type == 'text/csv':
            df = pd.read_csv(uploaded_file)
            content = '\n'.join(df['content'].tolist())
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            content = extract_text_from_docx(uploaded_file)
        else:  # txt file
            content = uploaded_file.read().decode()
        return content.strip()
    except Exception as e:
        st.error(f"Error reading file: {str(e)}")
        return ""

def process_with_chatgpt(content, prompt_type):
    try:
        prompt = PROMPTS[prompt_type].format(content=content)
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error processing with ChatGPT: {str(e)}")
        return ""

def main():
    st.title('Content Analysis App')
    
    # Add API Key input if not set in secrets
    if not API_KEY:
        api_key = st.text_input("Enter your OpenAI API Key:", type="password")
        if api_key:
            os.environ["OPENAI_API_KEY"] = api_key
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload your content file (TXT, CSV, or DOCX)", 
        type=['txt', 'csv', 'docx']
    )
    
    if uploaded_file is not None:
        # Process file content
        content = process_file_content(uploaded_file)
        
        if content:
            st.write("### Uploaded Content:")
            with st.expander("Show Content"):
                st.write(content)
            
            # Prompt selection
            prompt_type = st.selectbox(
                "Select Analysis Type",
                options=list(PROMPTS.keys()),
                format_func=lambda x: x.replace('_', ' ').title()
            )
            
            # Add process button
            if st.button("Analyze Content"):
                with st.spinner('Processing with ChatGPT...'):
                    output = process_with_chatgpt(content, prompt_type)
                    st.write("### Analysis Output:")
                    st.text_area("Generated Analysis", value=output, height=300)
                    
                    # Create and enable download button
                    if output:
                        doc = Document()
                        doc.add_heading('Original Content', level=1)
                        doc.add_paragraph(content)
                        doc.add_heading(f'Analysis ({prompt_type.replace("_", " ").title()})', level=1)
                        doc.add_paragraph(output)
                        
                        # Save the document to a bytes buffer
                        doc_buffer = io.BytesIO()
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        
                        st.download_button(
                            label="Download Analysis as Word Document",
                            data=doc_buffer,
                            file_name=f"{prompt_type}_analysis.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == '__main__':
    main() 